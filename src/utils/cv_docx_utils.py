from __future__ import annotations

from typing import Protocol

from docx import Document


class ParagraphUpdate(Protocol):
    index: int
    new_text: str


def extract_paragraphs(doc: Document) -> list[str]:
    return [paragraph.text for paragraph in doc.paragraphs]


def format_paragraphs_for_prompt(paragraphs: list[str]) -> str:
    lines: list[str] = []
    for i, text in enumerate(paragraphs):
        preview = text.replace("\n", " ").strip()
        if len(preview) > 180:
            preview = f"{preview[:177]}..."
        lines.append(f"[{i}] {preview}")
    return "\n".join(lines)


def apply_paragraph_changes(doc: Document, changes: list[ParagraphUpdate]) -> None:
    paragraphs = doc.paragraphs
    for change in changes:
        idx = change.index
        if 0 <= idx < len(paragraphs):
            paragraph = paragraphs[idx]
            if paragraph.runs:
                paragraph.runs[0].text = change.new_text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(change.new_text)
